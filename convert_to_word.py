#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文档转换为Word文档
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def set_heading_style(paragraph, level):
    """设置标题样式"""
    if level == 1:
        paragraph.style = 'Heading 1'
        paragraph.runs[0].font.size = Pt(18)
        paragraph.runs[0].font.bold = True
    elif level == 2:
        paragraph.style = 'Heading 2'
        paragraph.runs[0].font.size = Pt(16)
        paragraph.runs[0].font.bold = True
    elif level == 3:
        paragraph.style = 'Heading 3'
        paragraph.runs[0].font.size = Pt(14)
        paragraph.runs[0].font.bold = True
    elif level == 4:
        paragraph.style = 'Heading 4'
        paragraph.runs[0].font.size = Pt(12)
        paragraph.runs[0].font.bold = True

def add_code_block(doc, code_text, language=''):
    """添加代码块"""
    # 创建代码段落
    p = doc.add_paragraph()
    p.style = 'Normal'
    
    # 设置代码字体（等宽字体）
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    try:
        rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:cs'), 'Courier New')
    except:
        pass
    
    # 设置段落格式（缩进、背景色等）
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    return p

def convert_markdown_to_word(md_file, output_file, is_source_code=False):
    """将Markdown文件转换为Word文档"""
    print(f"正在转换: {md_file} -> {output_file}")
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    try:
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except:
        pass
    doc.styles['Normal'].font.size = Pt(12)
    
    # 设置代码字体
    try:
        code_style = doc.styles.add_style('Code', 1)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
    except:
        pass
    
    # 如果是源代码鉴别材料，添加页眉
    if is_source_code:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "数字生命体交互系统（心域）V1.0.0 源代码"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(10)
    
    # 解析Markdown内容
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_language = ''
    page_num = 1
    
    while i < len(lines):
        line = lines[i]
        
        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_block_lines:
                    code_text = '\n'.join(code_block_lines)
                    add_code_block(doc, code_text, code_language)
                    code_block_lines = []
                in_code_block = False
                code_language = ''
            else:
                # 开始代码块
                in_code_block = True
                code_language = line.strip()[3:].strip()
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            if title:
                p = doc.add_paragraph(title)
                set_heading_style(p, level)
                i += 1
                continue
        
        # 处理列表
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue
        
        # 处理有序列表
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            if text:
                p = doc.add_paragraph(text, style='List Number')
                i += 1
                continue
        
        # 处理粗体
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        line = re.sub(r'__(.+?)__', r'\1', line)
        
        # 处理普通文本
        if line.strip():
            # 检查是否是文件路径标记
            if line.strip().startswith('**文件路径：'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip().replace('**', ''))
                run.font.bold = True
                run.font.size = Pt(11)
            else:
                p = doc.add_paragraph(line.strip())
        else:
            # 空行
            doc.add_paragraph()
        
        i += 1
    
    # 处理最后的代码块
    if in_code_block and code_block_lines:
        code_text = '\n'.join(code_block_lines)
        add_code_block(doc, code_text, code_language)
    
    # 如果是源代码鉴别材料，添加页脚
    if is_source_code:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        footer_para._element.append(fldChar1)
        footer_para._element.append(instrText)
        footer_para._element.append(fldChar2)
    
    # 保存文档
    doc.save(output_file)
    print(f"✓ 已生成: {output_file}")

def main():
    """主函数"""
    base_dir = '/Users/admin/Documents/trae_projects/heartsphere_new'
    
    # 需要转换的文件列表
    files_to_convert = [
        ('软件说明书.md', '软件说明书.docx', False),
        ('源代码说明文档.md', '源代码说明文档.docx', False),
        ('源代码清单.md', '源代码清单.docx', False),
        ('软件著作权申请材料清单.md', '软件著作权申请材料清单.docx', False),
        ('源代码鉴别材料（前30页）.md', '源代码鉴别材料（前30页）.docx', True),
        ('源代码鉴别材料（后30页）.md', '源代码鉴别材料（后30页）.docx', True),
    ]
    
    for md_file, docx_file, is_source_code in files_to_convert:
        md_path = os.path.join(base_dir, md_file)
        docx_path = os.path.join(base_dir, docx_file)
        
        if os.path.exists(md_path):
            try:
                convert_markdown_to_word(md_path, docx_path, is_source_code)
            except Exception as e:
                print(f"✗ 转换失败 {md_file}: {e}")
        else:
            print(f"✗ 文件不存在: {md_file}")
    
    print("\n所有文档转换完成！")

if __name__ == '__main__':
    main()





